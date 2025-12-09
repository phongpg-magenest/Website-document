"""
Service for exporting generated documents to various formats (DOCX, PDF, MD, HTML)
"""
import io
import re
import logging
from typing import Optional
from enum import Enum

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import markdown

logger = logging.getLogger(__name__)


class ExportFormat(str, Enum):
    DOCX = "docx"
    PDF = "pdf"
    MD = "md"
    HTML = "html"


class ExportService:
    """Service for converting markdown content to various document formats"""

    def __init__(self):
        self._register_fonts()

    def _register_fonts(self):
        """Register fonts for PDF generation with Vietnamese support"""
        try:
            pdfmetrics.registerFont(TTFont('DejaVuSans', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
            pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
        except Exception as e:
            logger.warning(f"Could not register DejaVu fonts: {e}. Using default fonts.")

    async def export_document(
        self,
        content: str,
        format: ExportFormat,
        title: str = "Generated Document",
        document_type: str = "document"
    ) -> tuple[bytes, str, str]:
        """
        Export markdown content to specified format

        Returns:
            tuple: (file_bytes, filename, content_type)
        """
        if format == ExportFormat.DOCX:
            return await self._export_docx(content, title, document_type)
        elif format == ExportFormat.PDF:
            return await self._export_pdf(content, title, document_type)
        elif format == ExportFormat.MD:
            return await self._export_md(content, title, document_type)
        elif format == ExportFormat.HTML:
            return await self._export_html(content, title, document_type)
        else:
            raise ValueError(f"Unsupported export format: {format}")

    async def _export_md(self, content: str, title: str, document_type: str) -> tuple[bytes, str, str]:
        """Export as Markdown file"""
        filename = f"{self._sanitize_filename(title)}_{document_type}.md"
        return content.encode('utf-8'), filename, "text/markdown"

    async def _export_html(self, content: str, title: str, document_type: str) -> tuple[bytes, str, str]:
        """Export as HTML file"""
        html_content = markdown.markdown(
            content,
            extensions=['tables', 'fenced_code', 'toc']
        )

        full_html = f"""<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            color: #333;
        }}
        h1 {{ color: #2563eb; border-bottom: 2px solid #2563eb; padding-bottom: 10px; }}
        h2 {{ color: #1d4ed8; margin-top: 30px; }}
        h3 {{ color: #3b82f6; }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #f3f4f6;
            font-weight: bold;
        }}
        tr:nth-child(even) {{ background-color: #f9fafb; }}
        code {{
            background-color: #f3f4f6;
            padding: 2px 6px;
            border-radius: 4px;
            font-family: 'Courier New', monospace;
        }}
        pre {{
            background-color: #1f2937;
            color: #f9fafb;
            padding: 16px;
            border-radius: 8px;
            overflow-x: auto;
        }}
        pre code {{
            background-color: transparent;
            padding: 0;
        }}
        blockquote {{
            border-left: 4px solid #2563eb;
            margin: 20px 0;
            padding-left: 20px;
            color: #6b7280;
        }}
        ul, ol {{
            padding-left: 30px;
        }}
        li {{
            margin: 8px 0;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

        filename = f"{self._sanitize_filename(title)}_{document_type}.html"
        return full_html.encode('utf-8'), filename, "text/html"

    async def _export_docx(self, content: str, title: str, document_type: str) -> tuple[bytes, str, str]:
        """Export as DOCX file"""
        doc = Document()

        # Set up styles
        self._setup_docx_styles(doc)

        # Parse markdown and add to document
        lines = content.split('\n')
        i = 0
        in_code_block = False
        code_content = []
        in_table = False
        table_rows = []

        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    self._add_code_block(doc, '\n'.join(code_content))
                    code_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_content.append(line)
                i += 1
                continue

            # Handle tables
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_rows = []

                # Skip separator rows (|---|---|)
                if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                    i += 1
                    continue

                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells:
                    table_rows.append(cells)
                i += 1
                continue
            elif in_table:
                # End of table
                self._add_table(doc, table_rows)
                table_rows = []
                in_table = False

            # Handle headers
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:].strip(), level=4)
            # Handle bullet points
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, text)
            # Handle numbered lists
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                p = doc.add_paragraph(style='List Number')
                self._add_formatted_text(p, text)
            # Handle blockquotes
            elif line.strip().startswith('>'):
                text = line.strip()[1:].strip()
                p = doc.add_paragraph()
                p.style = 'Quote'
                self._add_formatted_text(p, text)
            # Handle regular paragraphs
            elif line.strip():
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)

            i += 1

        # Handle remaining table
        if in_table and table_rows:
            self._add_table(doc, table_rows)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{self._sanitize_filename(title)}_{document_type}.docx"
        return buffer.getvalue(), filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def _setup_docx_styles(self, doc: Document):
        """Setup custom styles for DOCX document"""
        styles = doc.styles

        # Quote style
        if 'Quote' not in [s.name for s in styles]:
            quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_style.font.italic = True
            quote_style.font.color.rgb = RGBColor(107, 114, 128)
            quote_style.paragraph_format.left_indent = Inches(0.5)

    def _add_formatted_text(self, paragraph, text: str):
        """Add text with inline formatting (bold, italic, code)"""
        # Split by formatting patterns
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            else:
                paragraph.add_run(part)

    def _add_code_block(self, doc: Document, code: str):
        """Add a code block to the document"""
        p = doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.25)

    def _add_table(self, doc: Document, rows: list):
        """Add a table to the document"""
        if not rows:
            return

        num_cols = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'

        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = row.cells[j]
                    cell.text = cell_text
                    # Bold the header row
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        doc.add_paragraph()  # Add spacing after table

    async def _export_pdf(self, content: str, title: str, document_type: str) -> tuple[bytes, str, str]:
        """Export as PDF file"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        # Setup styles
        styles = getSampleStyleSheet()

        # Try to use DejaVu font for Vietnamese support
        try:
            styles.add(ParagraphStyle(
                name='VietnameseNormal',
                fontName='DejaVuSans',
                fontSize=11,
                leading=14,
                spaceBefore=6,
                spaceAfter=6,
            ))
            styles.add(ParagraphStyle(
                name='VietnameseHeading1',
                fontName='DejaVuSans-Bold',
                fontSize=18,
                leading=22,
                spaceBefore=12,
                spaceAfter=6,
                textColor=colors.HexColor('#2563eb'),
            ))
            styles.add(ParagraphStyle(
                name='VietnameseHeading2',
                fontName='DejaVuSans-Bold',
                fontSize=14,
                leading=18,
                spaceBefore=10,
                spaceAfter=4,
                textColor=colors.HexColor('#1d4ed8'),
            ))
            styles.add(ParagraphStyle(
                name='VietnameseHeading3',
                fontName='DejaVuSans-Bold',
                fontSize=12,
                leading=16,
                spaceBefore=8,
                spaceAfter=4,
                textColor=colors.HexColor('#3b82f6'),
            ))
            normal_style = styles['VietnameseNormal']
            h1_style = styles['VietnameseHeading1']
            h2_style = styles['VietnameseHeading2']
            h3_style = styles['VietnameseHeading3']
        except Exception as e:
            logger.warning(f"Could not create Vietnamese styles: {e}. Using default.")
            normal_style = styles['Normal']
            h1_style = styles['Heading1']
            h2_style = styles['Heading2']
            h3_style = styles['Heading3']

        story = []

        # Parse markdown content
        lines = content.split('\n')
        i = 0
        in_code_block = False
        code_content = []
        in_table = False
        table_rows = []

        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    code_text = '\n'.join(code_content)
                    story.append(Paragraph(f"<pre>{code_text}</pre>", styles['Code'] if 'Code' in styles else normal_style))
                    code_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_content.append(line)
                i += 1
                continue

            # Handle tables
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_rows = []

                if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                    i += 1
                    continue

                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells:
                    table_rows.append(cells)
                i += 1
                continue
            elif in_table:
                self._add_pdf_table(story, table_rows, normal_style)
                table_rows = []
                in_table = False

            # Handle headers
            if line.startswith('# '):
                text = self._escape_pdf_text(line[2:].strip())
                story.append(Paragraph(text, h1_style))
            elif line.startswith('## '):
                text = self._escape_pdf_text(line[3:].strip())
                story.append(Paragraph(text, h2_style))
            elif line.startswith('### '):
                text = self._escape_pdf_text(line[4:].strip())
                story.append(Paragraph(text, h3_style))
            elif line.startswith('#### '):
                text = self._escape_pdf_text(line[5:].strip())
                story.append(Paragraph(f"<b>{text}</b>", normal_style))
            # Handle bullet points
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = self._escape_pdf_text(line.strip()[2:])
                text = self._convert_markdown_formatting(text)
                story.append(Paragraph(f"• {text}", normal_style))
            # Handle numbered lists
            elif re.match(r'^\d+\.\s', line.strip()):
                match = re.match(r'^(\d+)\.\s(.+)$', line.strip())
                if match:
                    num = match.group(1)
                    text = self._escape_pdf_text(match.group(2))
                    text = self._convert_markdown_formatting(text)
                    story.append(Paragraph(f"{num}. {text}", normal_style))
            # Handle blockquotes
            elif line.strip().startswith('>'):
                text = self._escape_pdf_text(line.strip()[1:].strip())
                story.append(Paragraph(f"<i>{text}</i>", normal_style))
            # Handle regular paragraphs
            elif line.strip():
                text = self._escape_pdf_text(line)
                text = self._convert_markdown_formatting(text)
                story.append(Paragraph(text, normal_style))
            else:
                story.append(Spacer(1, 6))

            i += 1

        # Handle remaining table
        if in_table and table_rows:
            self._add_pdf_table(story, table_rows, normal_style)

        doc.build(story)
        buffer.seek(0)

        filename = f"{self._sanitize_filename(title)}_{document_type}.pdf"
        return buffer.getvalue(), filename, "application/pdf"

    def _escape_pdf_text(self, text: str) -> str:
        """Escape special characters for PDF"""
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        return text

    def _convert_markdown_formatting(self, text: str) -> str:
        """Convert markdown bold/italic to PDF markup"""
        # Bold
        text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
        # Italic
        text = re.sub(r'\*([^*]+)\*', r'<i>\1</i>', text)
        # Code
        text = re.sub(r'`([^`]+)`', r'<font name="Courier">\1</font>', text)
        return text

    def _add_pdf_table(self, story: list, rows: list, style: ParagraphStyle):
        """Add a table to PDF"""
        if not rows:
            return

        # Create table data with Paragraphs for text wrapping
        table_data = []
        for row in rows:
            table_row = []
            for cell in row:
                cell_text = self._escape_pdf_text(cell)
                table_row.append(Paragraph(cell_text, style))
            table_data.append(table_row)

        if not table_data:
            return

        # Calculate column widths
        num_cols = max(len(row) for row in table_data)
        col_width = (A4[0] - 144) / num_cols  # Account for margins

        table = Table(table_data, colWidths=[col_width] * num_cols)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f4f6')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#111827')),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#d1d5db')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9fafb')]),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ]))

        story.append(table)
        story.append(Spacer(1, 12))

    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for safe file system usage"""
        # Remove or replace invalid characters
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)
        filename = filename.replace(' ', '_')
        # Limit length
        if len(filename) > 50:
            filename = filename[:50]
        return filename or "document"


# Singleton instance
export_service = ExportService()
