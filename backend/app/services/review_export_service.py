"""
Review Export Service - Export báo cáo review ra PDF và Word
"""
from typing import List
from io import BytesIO
from datetime import datetime
import logging

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, ListFlowable, ListItem
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.schemas.review import ReviewResult, SeverityLevel

logger = logging.getLogger(__name__)


class ReviewExportService:
    def __init__(self):
        # Try to register Vietnamese font for PDF
        try:
            # DejaVu supports Vietnamese
            pdfmetrics.registerFont(TTFont('DejaVu', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
            pdfmetrics.registerFont(TTFont('DejaVu-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
            self.font_name = 'DejaVu'
            self.font_bold = 'DejaVu-Bold'
        except:
            self.font_name = 'Helvetica'
            self.font_bold = 'Helvetica-Bold'

    def _get_severity_color(self, severity: SeverityLevel) -> colors.Color:
        """Get color for severity level"""
        severity_colors = {
            SeverityLevel.LOW: colors.Color(0.2, 0.6, 0.2),       # Green
            SeverityLevel.MEDIUM: colors.Color(0.9, 0.6, 0.1),    # Orange
            SeverityLevel.HIGH: colors.Color(0.8, 0.2, 0.2),      # Red
            SeverityLevel.CRITICAL: colors.Color(0.5, 0, 0.5),    # Purple
        }
        return severity_colors.get(severity, colors.black)

    def _get_severity_color_rgb(self, severity: SeverityLevel) -> RGBColor:
        """Get RGB color for severity level (for Word)"""
        severity_colors = {
            SeverityLevel.LOW: RGBColor(51, 153, 51),       # Green
            SeverityLevel.MEDIUM: RGBColor(230, 153, 26),   # Orange
            SeverityLevel.HIGH: RGBColor(204, 51, 51),      # Red
            SeverityLevel.CRITICAL: RGBColor(128, 0, 128),  # Purple
        }
        return severity_colors.get(severity, RGBColor(0, 0, 0))

    def _get_score_color(self, score: float) -> colors.Color:
        """Get color based on score"""
        if score >= 8:
            return colors.Color(0.2, 0.6, 0.2)  # Green
        elif score >= 6:
            return colors.Color(0.9, 0.6, 0.1)  # Orange
        else:
            return colors.Color(0.8, 0.2, 0.2)  # Red

    def export_to_pdf(self, review: ReviewResult) -> bytes:
        """Export review result to PDF"""
        buffer = BytesIO()

        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )

        # Styles
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=self.font_bold,
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER,
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=self.font_bold,
            fontSize=14,
            spaceBefore=15,
            spaceAfter=10,
            textColor=colors.Color(0.2, 0.3, 0.5),
        )

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=self.font_name,
            fontSize=10,
            spaceAfter=8,
        )

        # Build content
        story = []

        # Title
        story.append(Paragraph("BÁO CÁO REVIEW TÀI LIỆU", title_style))
        story.append(Spacer(1, 10))

        # Document info
        info_data = [
            ["Tên tài liệu:", review.document_name],
            ["Loại tài liệu:", review.document_type or "Không xác định"],
            ["Ngày review:", datetime.now().strftime("%d/%m/%Y %H:%M")],
            ["Thời gian xử lý:", f"{review.review_time_seconds:.1f} giây"],
        ]

        info_table = Table(info_data, colWidths=[4*cm, 12*cm])
        info_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), self.font_bold),
            ('FONTNAME', (1, 0), (1, -1), self.font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 20))

        # Overall Score
        story.append(Paragraph("ĐIỂM TỔNG THỂ", heading_style))

        score_color = self._get_score_color(review.overall_score)
        score_style = ParagraphStyle(
            'Score',
            parent=styles['Normal'],
            fontName=self.font_bold,
            fontSize=36,
            alignment=TA_CENTER,
            textColor=score_color,
        )
        story.append(Paragraph(f"{review.overall_score}/10", score_style))
        story.append(Spacer(1, 10))

        # Summary
        story.append(Paragraph(review.summary, normal_style))
        story.append(Spacer(1, 15))

        # Score breakdown
        story.append(Paragraph("ĐIỂM THEO HẠNG MỤC", heading_style))

        categories = review.categories
        score_data = [
            ["Hạng mục", "Điểm", "Số vấn đề"],
            [categories.spelling_grammar.label, f"{categories.spelling_grammar.score}/10", str(len(categories.spelling_grammar.issues))],
            [categories.structure.label, f"{categories.structure.score}/10", str(len(categories.structure.issues))],
            [categories.completeness.label, f"{categories.completeness.score}/10", str(len(categories.completeness.issues))],
            [categories.content_quality.label, f"{categories.content_quality.score}/10", str(len(categories.content_quality.issues))],
            [categories.risk_detection.label, f"{categories.risk_detection.score}/10", str(len(categories.risk_detection.risks))],
        ]

        score_table = Table(score_data, colWidths=[8*cm, 4*cm, 4*cm])
        score_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.2, 0.3, 0.5)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), self.font_bold),
            ('FONTNAME', (0, 1), (-1, -1), self.font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(score_table)
        story.append(Spacer(1, 20))

        # Detailed Issues
        story.append(Paragraph("CHI TIẾT CÁC VẤN ĐỀ", heading_style))

        # Function to add issues table
        def add_issues_section(title, issues, is_risk=False):
            if not issues:
                return

            story.append(Paragraph(f"<b>{title}</b>", normal_style))

            if is_risk:
                headers = ["Vị trí", "Rủi ro", "Tác động", "Mức độ"]
                data = [headers]
                for item in issues:
                    data.append([
                        item.location,
                        item.risk,
                        item.impact,
                        item.severity.value.upper()
                    ])
            else:
                headers = ["Vị trí", "Vấn đề", "Gợi ý", "Mức độ"]
                data = [headers]
                for item in issues:
                    data.append([
                        item.location,
                        item.issue,
                        item.suggestion or "-",
                        item.severity.value.upper()
                    ])

            if len(data) > 1:
                issue_table = Table(data, colWidths=[3*cm, 5*cm, 5*cm, 2*cm])
                issue_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.8, 0.8, 0.8)),
                    ('FONTNAME', (0, 0), (-1, 0), self.font_bold),
                    ('FONTNAME', (0, 1), (-1, -1), self.font_name),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                ]))
                story.append(issue_table)
                story.append(Spacer(1, 10))

        add_issues_section("Chính tả & Ngữ pháp", categories.spelling_grammar.issues)
        add_issues_section("Cấu trúc & Bố cục", categories.structure.issues)
        add_issues_section("Tính đầy đủ", categories.completeness.issues)
        add_issues_section("Chất lượng nội dung", categories.content_quality.issues)
        add_issues_section("Phát hiện rủi ro", categories.risk_detection.risks, is_risk=True)

        # Missing sections
        if categories.completeness.missing_sections:
            story.append(Paragraph("<b>Các mục bị thiếu:</b>", normal_style))
            for section in categories.completeness.missing_sections:
                story.append(Paragraph(f"• {section}", normal_style))
            story.append(Spacer(1, 10))

        # Recommendations
        if review.recommendations:
            story.append(Paragraph("KHUYẾN NGHỊ CẢI THIỆN", heading_style))
            for i, rec in enumerate(review.recommendations, 1):
                story.append(Paragraph(f"{i}. {rec}", normal_style))

        # Template comparison
        if review.template_comparison:
            story.append(Spacer(1, 15))
            story.append(Paragraph("SO SÁNH VỚI TEMPLATE", heading_style))
            story.append(Paragraph(f"Template: {review.template_comparison.template_name}", normal_style))

            if review.template_comparison.matched_sections:
                story.append(Paragraph("<b>Các mục khớp:</b> " + ", ".join(review.template_comparison.matched_sections), normal_style))

            if review.template_comparison.missing_sections:
                story.append(Paragraph("<b>Các mục thiếu:</b> " + ", ".join(review.template_comparison.missing_sections), normal_style))

            if review.template_comparison.extra_sections:
                story.append(Paragraph("<b>Các mục thừa:</b> " + ", ".join(review.template_comparison.extra_sections), normal_style))

        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def export_to_docx(self, review: ReviewResult) -> bytes:
        """Export review result to Word document"""
        doc = Document()

        # Title
        title = doc.add_heading("BÁO CÁO REVIEW TÀI LIỆU", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Document info
        doc.add_paragraph()
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'

        info_data = [
            ("Tên tài liệu:", review.document_name),
            ("Loại tài liệu:", review.document_type or "Không xác định"),
            ("Ngày review:", datetime.now().strftime("%d/%m/%Y %H:%M")),
            ("Thời gian xử lý:", f"{review.review_time_seconds:.1f} giây"),
        ]

        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
            info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # Overall Score
        doc.add_heading("ĐIỂM TỔNG THỂ", level=1)

        score_para = doc.add_paragraph()
        score_run = score_para.add_run(f"{review.overall_score}/10")
        score_run.bold = True
        score_run.font.size = Pt(36)
        if review.overall_score >= 8:
            score_run.font.color.rgb = RGBColor(51, 153, 51)
        elif review.overall_score >= 6:
            score_run.font.color.rgb = RGBColor(230, 153, 26)
        else:
            score_run.font.color.rgb = RGBColor(204, 51, 51)
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(review.summary)

        # Score breakdown
        doc.add_heading("ĐIỂM THEO HẠNG MỤC", level=1)

        categories = review.categories
        score_table = doc.add_table(rows=6, cols=3)
        score_table.style = 'Table Grid'

        headers = ["Hạng mục", "Điểm", "Số vấn đề"]
        for i, header in enumerate(headers):
            cell = score_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        scores_data = [
            (categories.spelling_grammar.label, categories.spelling_grammar.score, len(categories.spelling_grammar.issues)),
            (categories.structure.label, categories.structure.score, len(categories.structure.issues)),
            (categories.completeness.label, categories.completeness.score, len(categories.completeness.issues)),
            (categories.content_quality.label, categories.content_quality.score, len(categories.content_quality.issues)),
            (categories.risk_detection.label, categories.risk_detection.score, len(categories.risk_detection.risks)),
        ]

        for i, (label, score, count) in enumerate(scores_data, 1):
            score_table.rows[i].cells[0].text = label
            score_table.rows[i].cells[1].text = f"{score}/10"
            score_table.rows[i].cells[2].text = str(count)

        doc.add_paragraph()

        # Detailed Issues
        doc.add_heading("CHI TIẾT CÁC VẤN ĐỀ", level=1)

        def add_issues_section_docx(title, issues, is_risk=False):
            if not issues:
                return

            doc.add_heading(title, level=2)

            if is_risk:
                table = doc.add_table(rows=len(issues) + 1, cols=4)
                table.style = 'Table Grid'
                headers = ["Vị trí", "Rủi ro", "Tác động", "Mức độ"]
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header
                    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

                for i, item in enumerate(issues, 1):
                    table.rows[i].cells[0].text = item.location
                    table.rows[i].cells[1].text = item.risk
                    table.rows[i].cells[2].text = item.impact
                    severity_cell = table.rows[i].cells[3]
                    severity_cell.text = item.severity.value.upper()
                    severity_cell.paragraphs[0].runs[0].font.color.rgb = self._get_severity_color_rgb(item.severity)
            else:
                table = doc.add_table(rows=len(issues) + 1, cols=4)
                table.style = 'Table Grid'
                headers = ["Vị trí", "Vấn đề", "Gợi ý", "Mức độ"]
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header
                    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

                for i, item in enumerate(issues, 1):
                    table.rows[i].cells[0].text = item.location
                    table.rows[i].cells[1].text = item.issue
                    table.rows[i].cells[2].text = item.suggestion or "-"
                    severity_cell = table.rows[i].cells[3]
                    severity_cell.text = item.severity.value.upper()
                    severity_cell.paragraphs[0].runs[0].font.color.rgb = self._get_severity_color_rgb(item.severity)

            doc.add_paragraph()

        add_issues_section_docx("Chính tả & Ngữ pháp", categories.spelling_grammar.issues)
        add_issues_section_docx("Cấu trúc & Bố cục", categories.structure.issues)
        add_issues_section_docx("Tính đầy đủ", categories.completeness.issues)
        add_issues_section_docx("Chất lượng nội dung", categories.content_quality.issues)
        add_issues_section_docx("Phát hiện rủi ro", categories.risk_detection.risks, is_risk=True)

        # Missing sections
        if categories.completeness.missing_sections:
            doc.add_heading("Các mục bị thiếu", level=2)
            for section in categories.completeness.missing_sections:
                doc.add_paragraph(section, style='List Bullet')

        # Recommendations
        if review.recommendations:
            doc.add_heading("KHUYẾN NGHỊ CẢI THIỆN", level=1)
            for i, rec in enumerate(review.recommendations, 1):
                doc.add_paragraph(f"{i}. {rec}")

        # Template comparison
        if review.template_comparison:
            doc.add_heading("SO SÁNH VỚI TEMPLATE", level=1)
            doc.add_paragraph(f"Template: {review.template_comparison.template_name}")

            if review.template_comparison.matched_sections:
                p = doc.add_paragraph()
                p.add_run("Các mục khớp: ").bold = True
                p.add_run(", ".join(review.template_comparison.matched_sections))

            if review.template_comparison.missing_sections:
                p = doc.add_paragraph()
                p.add_run("Các mục thiếu: ").bold = True
                p.add_run(", ".join(review.template_comparison.missing_sections))

            if review.template_comparison.extra_sections:
                p = doc.add_paragraph()
                p.add_run("Các mục thừa: ").bold = True
                p.add_run(", ".join(review.template_comparison.extra_sections))

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


# Singleton instance
review_export_service = ReviewExportService()
