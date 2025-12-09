import os
import re
from typing import List, Optional, Tuple
from uuid import UUID
import logging
from io import BytesIO

from docx import Document as DocxDocument
from openpyxl import load_workbook
from PyPDF2 import PdfReader
import markdown
import tiktoken

from app.core.config import settings
from app.models.document import FileType

logger = logging.getLogger(__name__)


class DocumentProcessingService:
    def __init__(self):
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        self.chunk_size = 500  # tokens
        self.chunk_overlap = 100  # tokens

    def extract_text(self, file_content: bytes, file_type: FileType) -> str:
        """Extract text from different file formats"""
        try:
            if file_type in [FileType.DOC, FileType.DOCX]:
                return self._extract_docx(file_content)
            elif file_type in [FileType.XLS, FileType.XLSX]:
                return self._extract_xlsx(file_content)
            elif file_type == FileType.PDF:
                return self._extract_pdf(file_content)
            elif file_type == FileType.MD:
                return self._extract_markdown(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            raise

    def _extract_docx(self, file_content: bytes) -> str:
        """Extract text from Word documents"""
        doc = DocxDocument(BytesIO(file_content))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return "\n\n".join(paragraphs)

    def _extract_xlsx(self, file_content: bytes) -> str:
        """Extract text from Excel files"""
        wb = load_workbook(BytesIO(file_content), read_only=True, data_only=True)
        text_parts = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"## Sheet: {sheet_name}\n")

            rows = []
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(row_values):
                    rows.append(" | ".join(row_values))

            text_parts.extend(rows)

        wb.close()
        return "\n".join(text_parts)

    def _extract_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF files"""
        reader = PdfReader(BytesIO(file_content))
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)

        return "\n\n".join(text_parts)

    def _extract_markdown(self, file_content: bytes) -> str:
        """Extract text from Markdown files (remove formatting)"""
        text = file_content.decode("utf-8")
        # Convert markdown to HTML and then strip tags
        html = markdown.markdown(text)
        # Simple HTML tag removal
        clean_text = re.sub(r'<[^>]+>', '', html)
        return clean_text

    def chunk_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap"""
        tokens = self.tokenizer.encode(text)
        chunks = []

        start = 0
        while start < len(tokens):
            end = start + self.chunk_size
            chunk_tokens = tokens[start:end]
            chunk_text = self.tokenizer.decode(chunk_tokens)
            chunks.append(chunk_text)

            # Move start with overlap
            start = end - self.chunk_overlap
            if start >= len(tokens):
                break

        return chunks

    def get_file_type(self, filename: str) -> Optional[FileType]:
        """Get FileType enum from filename"""
        ext = os.path.splitext(filename)[1].lower()
        ext_mapping = {
            ".doc": FileType.DOC,
            ".docx": FileType.DOCX,
            ".xls": FileType.XLS,
            ".xlsx": FileType.XLSX,
            ".pdf": FileType.PDF,
            ".md": FileType.MD,
        }
        return ext_mapping.get(ext)

    def validate_file(self, filename: str, file_size: int) -> Tuple[bool, str]:
        """Validate file before processing"""
        # Check extension
        ext = os.path.splitext(filename)[1].lower()
        if ext not in settings.allowed_extensions_list:
            return False, f"File type {ext} is not supported. Allowed: {settings.ALLOWED_EXTENSIONS}"

        # Check size
        if file_size > settings.max_file_size_bytes:
            return False, f"File size exceeds maximum allowed ({settings.MAX_FILE_SIZE_MB}MB)"

        return True, "OK"

    def generate_snippet(self, text: str, query: str, max_length: int = 300) -> str:
        """Generate a snippet with highlighted query terms"""
        # Find the position of query terms in text
        query_terms = query.lower().split()
        text_lower = text.lower()

        best_pos = 0
        best_score = 0

        for i in range(0, len(text) - max_length, 50):
            segment = text_lower[i:i + max_length]
            score = sum(1 for term in query_terms if term in segment)
            if score > best_score:
                best_score = score
                best_pos = i

        snippet = text[best_pos:best_pos + max_length]

        # Add ellipsis if needed
        if best_pos > 0:
            snippet = "..." + snippet
        if best_pos + max_length < len(text):
            snippet = snippet + "..."

        return snippet

    def highlight_terms(self, text: str, query: str) -> List[str]:
        """Extract highlighted passages containing query terms"""
        query_terms = query.lower().split()
        sentences = text.split(".")
        highlights = []

        for sentence in sentences:
            sentence_lower = sentence.lower()
            if any(term in sentence_lower for term in query_terms):
                highlights.append(sentence.strip() + ".")
                if len(highlights) >= 3:
                    break

        return highlights


# Singleton instance
document_processing_service = DocumentProcessingService()
